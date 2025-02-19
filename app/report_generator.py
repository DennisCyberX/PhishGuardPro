from docx import Document

def generate_report(user_input, analysis_result):
    document = Document()
    document.add_heading('PhishGuardPro Analysis Report', level=1)
    document.add_paragraph(f'Input: {user_input}')
    document.add_paragraph(f'Is Phishing: {analysis_result["is_phishing"]}')
    document.add_paragraph(f'Details: {analysis_result["details"]}')
    report_path = 'report.docx'
    document.save(report_path)
    return report_path
